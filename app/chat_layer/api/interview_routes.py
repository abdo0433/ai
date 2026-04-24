"""
AI Interview Platform + ATS Analyzer - FastAPI Backend v2.5
============================================================
Merged from:
  - AI Interview Platform v2.5 (session management, questions, evaluation, audio/video analysis)
  - ATS Analyzer (CV vs Job Description matching + AI feedback via Groq)

Run: uvicorn main:app --reload --port 8000
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Dict, List, Optional
from collections import Counter
from datetime import datetime
import io, os, json, uuid, time, tempfile
import soundfile as sf
from app.chat_layer.config import GROQ_API_KEY,GROQ_API_KEY2

# ============================================================
# ENGINE IMPORTS  (Interview Platform)
# ============================================================

print("\n" + "=" * 60)
print("🔄 Loading AI Interview Engines...")
print("=" * 60)

LLM_CLIENT = None
try:
    from app.chat_layer.core.llm_client import call_llm
    LLM_CLIENT = call_llm
    print("✅ LLM Client (Groq) loaded")
except Exception as e:
    print(f"⚠️  LLM Client: {e}")
    def call_llm(prompt: str) -> str:
        return '{"error": "LLM not available"}'
    LLM_CLIENT = call_llm

AUDIO_ENGINE = None
try:
    from app.chat_layer.engines.audio_feature import AudioFeatureExtractor
    AUDIO_ENGINE = AudioFeatureExtractor()
    print("✅ Audio Engine loaded")
except Exception as e:
    print(f"⚠️  Audio Engine: {e}")

VIDEO_ENGINE  = None
VIDEO_VERSION = "None"
try:
    from app.chat_layer.engines.video_feature_engine import VideoFeatureExtractor
    _ve = VideoFeatureExtractor()
    VIDEO_VERSION = "V2.0 Enhanced" if hasattr(_ve, "alert_manager") else "V1.0"
    VIDEO_ENGINE  = _ve
    print(f"✅ Video Engine {VIDEO_VERSION} loaded")
except Exception as e:
    print(f"⚠️  Video Engine: {e}")

EVAL_FN = None
try:
    from app.chat_layer.engines.evaluation_engine_multimodel import evaluate_answer as _eval_fn
    EVAL_FN = _eval_fn
    print("✅ Evaluation Engine loaded")
except Exception as e:
    print(f"⚠️  Evaluation Engine: {e}")

QUESTION_FN = None
try:
    from app.chat_layer.engines.question_engine import generate_question as _q_fn
    QUESTION_FN = _q_fn
    print("✅ Question Engine loaded")
except Exception as e:
    print(f"⚠️  Question Engine: {e}")

STT_ENGINE = None
try:
    from app.chat_layer.engines.stt_engine import get_stt_engine
    STT_ENGINE = get_stt_engine(model_path="small", language="en")
    print("✅ STT Engine (Whisper) loaded")
except Exception as e:
    print(f"⚠️  STT Engine: {e}")

TTS_FN = None
try:
    from app.chat_layer.engines.tts_engine import generate_tts_base64
    TTS_FN = generate_tts_base64
    print("✅ TTS Engine loaded")
except Exception as e:
    print(f"⚠️  TTS Engine: {e}")
    def generate_tts_base64(text: str) -> Optional[str]:
        return None
    TTS_FN = generate_tts_base64

print("=" * 60 + "\n")

# ============================================================
# GROQ CLIENT  (ATS Analyzer)
# ============================================================



try:
    from groq import Groq
    groq_client = Groq(api_key=GROQ_API_KEY)
    print("✅ Groq client (ATS) loaded")
except Exception as e:
    groq_client = None
    print(f"⚠️  Groq client (ATS): {e}")

# ============================================================
# CV / DOCUMENT PARSER  (shared by both systems)
# ============================================================

_PYPDF_OK = False
_DOCX_OK  = False

try:
    from pypdf import PdfReader
    _PYPDF_OK = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        _PYPDF_OK = True
    except ImportError:
        pass

try:
    from docx import Document as DocxDocument
    _DOCX_OK = True
except ImportError:
    pass


def _parse_pdf_bytes(file_bytes: bytes) -> str:
    if not _PYPDF_OK:
        return "[CV_PARSE_ERROR] pypdf not installed. Run: pip install pypdf"
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages  = [p.extract_text() for p in reader.pages if p.extract_text()]
        text   = "\n\n".join(p.strip() for p in pages).strip()
        return text if text else "[CV_PARSE_ERROR] PDF has no extractable text."
    except Exception as exc:
        return f"[CV_PARSE_ERROR] PDF parse failed: {exc}"


def _parse_docx_bytes(file_bytes: bytes) -> str:
    if not _DOCX_OK:
        return "[CV_PARSE_ERROR] python-docx not installed. Run: pip install python-docx"
    try:
        doc   = DocxDocument(io.BytesIO(file_bytes))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in lines:
                        lines.append(t)
        text = "\n".join(lines).strip()
        return text if text else "[CV_PARSE_ERROR] DOCX appears empty."
    except Exception as exc:
        return f"[CV_PARSE_ERROR] DOCX parse failed: {exc}"


def _parse_doc_legacy(file_bytes: bytes) -> str:
    try:
        import subprocess
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        result = subprocess.run(
            ["pandoc", tmp_path, "-t", "plain"],
            capture_output=True, text=True, timeout=15,
        )
        os.remove(tmp_path)
        if result.returncode != 0:
            return f"[CV_PARSE_ERROR] pandoc error: {result.stderr.strip()}"
        text = result.stdout.strip()
        return text if text else "[CV_PARSE_ERROR] .doc produced no text."
    except FileNotFoundError:
        return "[CV_PARSE_ERROR] pandoc not installed."
    except Exception as exc:
        return f"[CV_PARSE_ERROR] .doc parse failed: {exc}"


def extract_cv_text(file_bytes: bytes, filename: str) -> str:
    """Universal CV parser — supports PDF, DOCX, DOC, TXT, MD."""
    ext = os.path.splitext(filename.lower())[1]
    if   ext == ".pdf":  return _parse_pdf_bytes(file_bytes)
    elif ext == ".docx": return _parse_docx_bytes(file_bytes)
    elif ext == ".doc":  return _parse_doc_legacy(file_bytes)
    elif ext in (".txt", ".md"):
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return file_bytes.decode(enc).strip()
            except UnicodeDecodeError:
                continue
        return "[CV_PARSE_ERROR] Could not decode text file."
    else:
        return f"[CV_PARSE_ERROR] Unsupported format '{ext}'."

# ============================================================
# ATS HELPERS
# ============================================================

def _ats_score(cv_text: str, jd_text: str):
    """
    Simple keyword-based ATS scoring.
    Returns (score_int, matched_keywords, missing_keywords).
    """
    import re
    def tokenize(text):
        return set(re.findall(r'\b[a-zA-Z]{3,}\b', text.lower()))

    cv_words = tokenize(cv_text)
    jd_words = tokenize(jd_text)

    # Remove very common stop-words
    stopwords = {
        "the", "and", "for", "with", "this", "that", "are", "you", "your",
        "our", "will", "have", "has", "from", "not", "but", "can", "was",
        "they", "their", "been", "more", "also", "about", "its", "into",
    }
    jd_keywords = jd_words - stopwords

    matched = list(cv_words & jd_keywords)
    missing = list(jd_keywords - cv_words)
    score   = round(len(matched) / len(jd_keywords) * 100) if jd_keywords else 0
    return min(score, 100), sorted(matched), sorted(missing)


def _llm_feedback(cv_text: str, jd_text: str, ats: int) -> str:
    """Call Groq LLM to generate detailed ATS feedback."""
    if not groq_client:
        return "Groq client not available. Please check your API key."

    prompt = f"""
You are a professional ATS resume analyzer.

Candidate CV:
{cv_text}

Job Description:
{jd_text}

ATS Score: {ats}%

Give a clear analysis:

1- What is good in the CV

2- Mistakes in the CV

3- Missing skills

4- How to improve the CV

5- Suggest better words to replace weak words

Write in simple clear sentences.
No bullet symbols.
"""
    try:
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=1024,
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Groq Error: {e}"

# ============================================================
# APP SETUP
# ============================================================

app = FastAPI(
    title="AI Interview Platform + ATS Analyzer API",
    description="Combined backend: AI-powered interview simulation & ATS CV scoring",
    version="2.5.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://localhost:5173",
        "http://localhost:8080",
        "*",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================
# CONSTANTS  (Interview Platform)
# ============================================================

STAGE_ORDER = ["introduction", "technical", "behavioral"]

STAGE_QUESTION_COUNT: Dict[str, int] = {
    "introduction": 3,
    "technical":    4,
    "behavioral":   3,
}

STAGE_FALLBACK_QUESTIONS: Dict[str, List[str]] = {
    "introduction": [
        "Tell me about yourself and your background.",
        "What motivated you to apply for this position?",
        "Where do you see yourself in 5 years?",
        "What are your key strengths?",
        "How would your friends describe you professionally?",
    ],
    "technical": [
        "Describe a challenging technical problem you solved recently.",
        "What technologies are you most comfortable with and why?",
        "How do you approach debugging a complex issue?",
        "Explain a concept from your field to a non-technical person.",
        "How do you keep your technical skills up to date?",
        "Describe your experience with version control systems.",
        "What is your process for reviewing or writing code?",
    ],
    "behavioral": [
        "Tell me about a time you worked under a tight deadline.",
        "Describe a conflict with a teammate and how you resolved it.",
        "Give an example of when you showed leadership.",
        "Tell me about a time you failed and what you learned.",
        "How do you handle receiving critical feedback?",
    ],
}

# ============================================================
# SESSION STORE
# ============================================================

sessions: Dict[str, Dict] = {}

# ============================================================
# PYDANTIC MODELS
# ============================================================

class SessionStartResponse(BaseModel):
    session_id:      str
    stage:           str
    stage_index:     int
    question_number: int
    total_questions: int
    cv_parsed:       bool
    status:          str
    message:         str

class NextQuestionResponse(BaseModel):
    question:        str
    question_audio:  Optional[str] = None
    stage:           str
    stage_index:     int
    question_number: int
    total_questions: int
    is_first:        bool = False
    message:         str  = ""

class AnswerSubmitRequest(BaseModel):
    session_id:     str
    answer:         str
    audio_features: Optional[Dict] = None
    video_features: Optional[Dict] = None

class AnswerSubmitResponse(BaseModel):
    success:             bool
    evaluation_summary:  Dict
    stage_complete:      bool = False
    interview_complete:  bool = False
    next_stage:          Optional[str] = None
    next_question:       Optional[str] = None
    next_question_audio: Optional[str] = None
    message:             str  = ""

class FinalFeedbackResponse(BaseModel):
    success:                 bool
    overall_score:           int
    technical_score:         int
    communication_score:     int
    confidence_index:        int
    stress_index:            int
    professional_presence:   int
    strengths:               List[str]
    weaknesses:              List[str]
    improvement_suggestions: List[str]
    stage_summaries:         Dict
    recommendation:          str
    message:                 str

class VideoFrameRequest(BaseModel):
    session_id:   str
    frame_base64: str

class QuickAnalysisRequest(BaseModel):
    question:       str
    answer:         str
    audio_features: Optional[Dict] = None
    video_features: Optional[Dict] = None

# ============================================================
# SESSION HELPERS
# ============================================================

def create_session(job_title: str, job_description: str, cv_text: str) -> Dict:
    return {
        "id":                     str(uuid.uuid4()),
        "job_title":              job_title,
        "job_description":        job_description,
        "cv_text":                cv_text,
        "created_at":             datetime.utcnow().isoformat(),
        "current_stage_index":    0,
        "current_stage":          STAGE_ORDER[0],
        "current_question_index": 0,
        "global_question_index":  0,
        "stages": {
            stage: {
                "questions":      [],
                "answers":        [],
                "evaluations":    [],
                "audio_features": [],
                "video_features": [],
            }
            for stage in STAGE_ORDER
        },
        "status":              "active",
        "video_frames_buffer": [],
    }


def _get_current_state(session: Dict) -> Dict:
    stage     = session["current_stage"]
    s_idx     = session["current_stage_index"]
    q_idx     = session["current_question_index"]
    max_q     = STAGE_QUESTION_COUNT[stage]
    answered  = {s: len(session["stages"][s]["answers"]) for s in STAGE_ORDER}
    total_ans = sum(answered.values())
    return {
        "session_id":             session["id"],
        "status":                 session["status"],
        "job_title":              session["job_title"],
        "cv_available":           bool(session.get("cv_text")),
        "current_stage":          stage,
        "stage_index":            s_idx,
        "current_question_index": q_idx,
        "question_number":        q_idx + 1,
        "total_questions":        max_q,
        "questions_answered":     answered,
        "total_answered":         total_ans,
        "global_total":           sum(STAGE_QUESTION_COUNT.values()),
        "created_at":             session["created_at"],
    }


def _generate_adaptive_question(session: Dict, stage: str) -> str:
    stage_data         = session["stages"][stage]
    previous_answers   = stage_data["answers"]
    previous_questions = stage_data["questions"]
    q_num              = len(previous_questions) + 1

    if QUESTION_FN:
        try:
            ctx  = f"Job Title: {session['job_title']}\n"
            ctx += f"Job Description: {session['job_description']}\n"
            ctx += f"CV: {session['cv_text']}\n"
            ctx += f"Interview Stage: {stage}\n"
            ctx += f"Question Number: {q_num}\n"
            if previous_answers:
                ctx += "\nPrevious Q&A in this stage:\n"
                for i, (pq, pa) in enumerate(zip(previous_questions, previous_answers)):
                    ctx += f"Q{i+1}: {pq}\nA{i+1}: {pa}\n\n"
                ctx += (
                    "Based on the above answers, generate the NEXT question. "
                    "If the candidate answered well, go deeper. "
                    "If they struggled, ask a simpler related question."
                )
            q = QUESTION_FN(session["job_title"], ctx, session["cv_text"], stage)
            if q and len(q.strip()) > 10:
                return q.strip()
        except Exception as e:
            print(f"⚠️  Adaptive question gen failed: {e}")

    import random
    templates = STAGE_FALLBACK_QUESTIONS.get(stage, [])
    used      = set(previous_questions)
    available = [q for q in templates if q not in used]
    return random.choice(available) if available else (
        random.choice(templates) if templates else "Tell me more about your experience."
    )


def _evaluate_answer_safe(
    question:       str,
    answer:         str,
    audio_features: Optional[Dict],
    video_features: Optional[Dict],
    job_title:      str,
) -> Dict:
    if EVAL_FN:
        try:
            result = EVAL_FN(
                question=question,
                answer=answer,
                audio_features=audio_features,
                video_features=video_features,
            )
            if "error" not in result:
                return result
        except Exception as e:
            print(f"⚠️  Evaluation error: {e}")

    word_count             = len(answer.split())
    base                   = min(100, 50 + (word_count // 4))
    conf, stress, presence = 70, 30, 70

    if audio_features:
        if audio_features.get("energy_mean", 0.15) > 0.15:
            conf += 8
        if 100 <= audio_features.get("speaking_rate_estimate", 130) <= 180:
            presence += 5

    if video_features and "behavioral" in video_features:
        beh    = video_features["behavioral"]
        conf   = int(beh.get("confidence_score", conf))
        stress = int(beh.get("stress_level",      stress))

    return {
        "technical_score":             base,
        "communication_score":         min(100, base + 10),
        "relevance_score":             base,
        "confidence_index":            min(100, conf),
        "stress_index":                max(0,   stress),
        "professional_presence_index": min(100, presence),
        "summary":                     f"{'Detailed' if word_count > 80 else 'Brief'} answer provided.",
        "short_feedback":              "Try to provide more specific examples." if word_count < 50 else "Well-structured answer.",
        "delivery_feedback":           "Work on maintaining consistent energy and pace.",
        "strengths":                   ["Clear communication"] if word_count > 40 else ["Attempting to answer"],
        "weaknesses":                  ["Could provide more detail"] if word_count < 50 else [],
        "overall_recommendation":      "Moderate Fit",
        "recommendation_justification":"Based on the answer provided.",
    }


def _compute_final_feedback(session: Dict) -> Dict:
    all_evals, stage_summaries = [], {}

    for stage in STAGE_ORDER:
        evals = session["stages"][stage]["evaluations"]
        if not evals:
            continue
        tech = [e.get("technical_score",    0)  for e in evals]
        comm = [e.get("communication_score", 0)  for e in evals]
        conf = [e.get("confidence_index",    70) for e in evals]
        stage_summaries[stage] = {
            "avg_technical":      round(sum(tech) / len(tech)),
            "avg_communication":  round(sum(comm) / len(comm)),
            "avg_confidence":     round(sum(conf) / len(conf)),
            "questions_answered": len(evals),
        }
        all_evals.extend(evals)

    if not all_evals:
        return {
            "overall_score": 0, "technical_score": 0, "communication_score": 0,
            "confidence_index": 0, "stress_index": 50, "professional_presence": 0,
            "strengths": [], "weaknesses": ["No answers recorded"],
            "improvement_suggestions": ["Complete the interview"],
            "stage_summaries": {}, "recommendation": "Incomplete",
        }

    def avg(key, default=0):
        vals = [e.get(key, default) for e in all_evals]
        return round(sum(vals) / len(vals)) if vals else default

    overall     = avg("technical_score") // 2 + avg("communication_score") // 4 + avg("confidence_index") // 4
    strengths   = ["Communication skills", "Confidence"]
    weaknesses  = ["Technical depth could be improved"]
    suggestions = ["Practice more technical scenarios"]

    if LLM_CLIENT:
        try:
            prompt = f"""Return ONLY valid JSON. No markdown.
Interview summaries for {session['job_title']}:
{json.dumps(stage_summaries, indent=2)}
Required:
{{"strengths":["up to 5"],"weaknesses":["up to 5"],"improvement_suggestions":["up to 5"]}}"""
            raw        = LLM_CLIENT(prompt)
            s, e       = raw.find("{"), raw.rfind("}") + 1
            if s != -1 and e > s:
                parsed      = json.loads(raw[s:e])
                strengths   = parsed.get("strengths",               strengths)
                weaknesses  = parsed.get("weaknesses",              weaknesses)
                suggestions = parsed.get("improvement_suggestions", suggestions)
        except Exception as ex:
            print(f"⚠️  LLM final feedback error: {ex}")

    overall_score = min(100, overall)
    return {
        "overall_score":           overall_score,
        "technical_score":         avg("technical_score"),
        "communication_score":     avg("communication_score"),
        "confidence_index":        avg("confidence_index"),
        "stress_index":            avg("stress_index", 30),
        "professional_presence":   avg("professional_presence_index", 70),
        "strengths":               strengths,
        "weaknesses":              weaknesses,
        "improvement_suggestions": suggestions,
        "stage_summaries":         stage_summaries,
        "recommendation": (
            "Strong Fit"       if overall_score >= 80 else
            "Moderate Fit"     if overall_score >= 60 else
            "Needs Improvement"
        ),
    }

# ============================================================
# ANALYSIS PAGE HELPERS
# ============================================================

def _aggregate_audio(session: Dict) -> Dict:
    all_audio = []
    for stage in STAGE_ORDER:
        all_audio.extend(session["stages"][stage]["audio_features"])
    if not all_audio:
        return {}

    def avg_f(key, default=0.0):
        vals = [a.get(key, default) for a in all_audio if isinstance(a.get(key), (int, float))]
        return round(sum(vals) / len(vals), 3) if vals else default

    return {
        "pitch_mean":             avg_f("pitch_mean"),
        "pitch_variability":      avg_f("pitch_variability"),
        "energy_mean":            avg_f("energy_mean"),
        "speaking_rate_estimate": avg_f("speaking_rate_estimate"),
        "silence_ratio":          avg_f("silence_ratio"),
        "num_pauses":             round(avg_f("num_pauses")),
        "avg_pause_duration":     avg_f("avg_pause_duration"),
        "overall_audio_quality":  avg_f("overall_audio_quality"),
        "samples":                len(all_audio),
    }


def _aggregate_video(session: Dict) -> Dict:
    all_video = []
    for stage in STAGE_ORDER:
        all_video.extend(session["stages"][stage]["video_features"])
    if not all_video:
        return {}

    def avg_n(key, subkey, default=0.0):
        vals = [
            v.get(key, {}).get(subkey, default)
            for v in all_video
            if isinstance(v.get(key, {}).get(subkey), (int, float))
        ]
        return round(sum(vals) / len(vals), 2) if vals else default

    def most_str(key, subkey, default="moderate"):
        vals = [v.get(key, {}).get(subkey, default) for v in all_video
                if isinstance(v.get(key, {}).get(subkey), str)]
        return Counter(vals).most_common(1)[0][0] if vals else default

    def any_bool(key, subkey):
        return any(v.get(key, {}).get(subkey, False) for v in all_video)

    emotion_keys = ["happy", "neutral", "surprised", "fear", "sad", "angry", "disgust"]
    emotions_avg = {k: avg_n("emotions", k) for k in emotion_keys}
    emotions_avg["positive_ratio"]   = avg_n("emotions", "positive_ratio")
    emotions_avg["negative_ratio"]   = avg_n("emotions", "negative_ratio")
    dom = [v.get("emotions", {}).get("dominant_emotion", "neutral") for v in all_video]
    emotions_avg["dominant_emotion"] = Counter(dom).most_common(1)[0][0] if dom else "neutral"

    return {
        "emotions": emotions_avg,
        "gaze": {
            "eye_contact_percentage":  avg_n("gaze", "eye_contact_percentage"),
            "looking_away_percentage": avg_n("gaze", "looking_away_percentage"),
            "gaze_stability":          avg_n("gaze", "gaze_stability"),
        },
        "behavioral": {
            "confidence_score":  avg_n("behavioral", "confidence_score"),
            "stress_level":      avg_n("behavioral", "stress_level"),
            "smile_percentage":  avg_n("behavioral", "smile_percentage"),
            "head_stability":    avg_n("behavioral", "head_stability"),
            "overall_composure": avg_n("behavioral", "overall_composure"),
        },
        "posture": {
            "average_posture_score": avg_n("posture", "average_posture_score"),
            "slouching_percentage":  avg_n("posture", "slouching_percentage"),
            "upright_percentage":    avg_n("posture", "upright_percentage"),
        },
        "eye_contact": {
            "eye_contact_percentage": avg_n("eye_contact", "eye_contact_percentage"),
            "eye_contact_quality":    most_str("eye_contact", "eye_contact_quality"),
            "engagement_level":       most_str("eye_contact", "engagement_level"),
        },
        "gestures": {
            "fidgeting_detected":  any_bool("gestures", "fidgeting_detected"),
            "fidgeting_severity":  most_str("gestures",  "fidgeting_severity", "none"),
            "expressive_gestures": any_bool("gestures", "expressive_gestures"),
        },
        "samples": len(all_video),
    }


def _score_communication(audio: Dict, video: Dict) -> Dict:
    energy        = audio.get("energy_mean", 0.1)
    audio_quality = audio.get("overall_audio_quality", 50)
    if   energy < 0.05: vocal_clarity, vocal_label = 30, "Too quiet"
    elif energy < 0.10: vocal_clarity, vocal_label = 55, "Slightly quiet"
    elif energy < 0.35: vocal_clarity, vocal_label = 85, "Good"
    else:               vocal_clarity, vocal_label = 70, "Loud"
    vocal_clarity = round((vocal_clarity + audio_quality) / 2)

    rate = audio.get("speaking_rate_estimate", 120)
    if   rate < 80:   pace_score, pace_label = 40, "Too slow"
    elif rate < 110:  pace_score, pace_label = 65, "Slightly slow"
    elif rate <= 180: pace_score, pace_label = 90, "Ideal"
    elif rate <= 220: pace_score, pace_label = 70, "Slightly fast"
    else:             pace_score, pace_label = 45, "Too fast"

    pv = audio.get("pitch_variability", 0.1)
    if   pv < 0.04: tone_score, tone_label = 35, "Monotone"
    elif pv < 0.10: tone_score, tone_label = 65, "Slightly varied"
    elif pv < 0.22: tone_score, tone_label = 90, "Well varied"
    else:           tone_score, tone_label = 70, "Very expressive"

    silence   = audio.get("silence_ratio", 0.2)
    avg_pause = audio.get("avg_pause_duration", 0.5)
    if   silence < 0.08:  pause_score, pause_label = 55, "Few pauses — rushing"
    elif silence <= 0.30: pause_score, pause_label = 90, "Natural pauses"
    elif silence <= 0.45: pause_score, pause_label = 65, "Many pauses"
    else:                 pause_score, pause_label = 40, "Excessive silence"
    if avg_pause > 2.5:
        pause_score  = max(30, pause_score - 15)
        pause_label += " (long gaps)"

    eye_pct = (
        video.get("gaze",        {}).get("eye_contact_percentage") or
        video.get("eye_contact", {}).get("eye_contact_percentage") or 50
    )
    if   eye_pct >= 70: eye_score, eye_label = 90, "Excellent"
    elif eye_pct >= 50: eye_score, eye_label = 70, "Good"
    elif eye_pct >= 30: eye_score, eye_label = 50, "Fair"
    else:               eye_score, eye_label = 30, "Needs improvement"

    pos_ratio  = video.get("emotions", {}).get("positive_ratio", 0.2)
    engagement = round(min(100, 40 + pos_ratio * 150))

    overall = round(
        vocal_clarity * 0.20 + pace_score * 0.20 + tone_score  * 0.20 +
        pause_score   * 0.15 + eye_score  * 0.15 + engagement  * 0.10
    )
    return {
        "overall_score": overall,
        "metrics": {
            "vocal_clarity":    {"score": vocal_clarity, "label": vocal_label,  "value": round(energy, 3)},
            "speaking_pace":    {"score": pace_score,    "label": pace_label,   "value": round(rate, 1), "unit": "syllables/min"},
            "tone_variation":   {"score": tone_score,    "label": tone_label,   "value": round(pv, 3)},
            "pause_management": {"score": pause_score,   "label": pause_label,  "value": round(silence, 2), "avg_pause_sec": round(avg_pause, 2)},
            "eye_contact":      {"score": eye_score,     "label": eye_label,    "value": round(eye_pct, 1), "unit": "%"},
            "engagement":       {"score": engagement,    "label": "Emotional presence", "value": round(pos_ratio, 2)},
        },
        "raw_audio": audio,
    }


def _score_body_language(video: Dict) -> Dict:
    if not video:
        return {"overall_score": 0, "metrics": {}, "no_data": True}

    beh      = video.get("behavioral", {})
    posture  = video.get("posture",    {})
    gestures = video.get("gestures",   {})
    gaze     = video.get("gaze",       {})
    emotions = video.get("emotions",   {})

    conf            = beh.get("confidence_score", 50)
    conf_label      = "High" if conf >= 70 else ("Moderate" if conf >= 45 else "Low")
    stress          = beh.get("stress_level", 40)
    stress_label    = "Low" if stress <= 30 else ("Moderate" if stress <= 55 else "High")
    composure_score = round(100 - stress)
    upright         = posture.get("upright_percentage", 75)
    posture_score   = round(posture.get("average_posture_score", 7.5) * 10)
    posture_label   = "Excellent" if upright >= 80 else ("Good" if upright >= 60 else "Fair")
    stability       = gaze.get("gaze_stability", 0.7)
    gaze_score      = round(stability * 100)
    gaze_label      = "Stable" if stability >= 0.75 else ("Moderate" if stability >= 0.5 else "Unstable")
    severity        = gestures.get("fidgeting_severity", "none")
    fidget_score    = {"none": 95, "low": 75, "moderate": 50, "high": 25}.get(severity, 75)
    smile_pct       = beh.get("smile_percentage", 20)
    expr            = gestures.get("expressive_gestures", True)
    smile_score     = min(100, round(40 + smile_pct * 1.5) + (10 if expr else 0))
    neg_ratio       = emotions.get("negative_ratio", 0.1)
    dom             = emotions.get("dominant_emotion", "neutral")
    emotion_score   = round(max(0, 100 - neg_ratio * 200))

    overall = round(
        conf           * 0.25 + composure_score * 0.20 +
        posture_score  * 0.15 + gaze_score       * 0.15 +
        fidget_score   * 0.10 + smile_score      * 0.10 +
        emotion_score  * 0.05
    )
    return {
        "overall_score": min(100, overall),
        "metrics": {
            "confidence":     {"score": round(conf),     "label": conf_label,               "value": round(conf, 1)},
            "composure":      {"score": composure_score, "label": f"{stress_label} stress",  "value": round(stress, 1), "note": "Lower stress = higher composure"},
            "posture":        {"score": posture_score,   "label": posture_label,             "value": round(upright, 1), "unit": "% upright"},
            "gaze_stability": {"score": gaze_score,      "label": gaze_label,                "value": round(stability, 3)},
            "fidgeting":      {"score": fidget_score,    "label": f"{severity.capitalize()} fidgeting", "value": gestures.get("fidgeting_detected", False)},
            "expressiveness": {"score": smile_score,     "label": "Smile & gestures",        "value": round(smile_pct, 1), "unit": "%"},
            "emotion_control":{"score": emotion_score,   "label": f"Dominant: {dom}",        "value": round(neg_ratio, 2)},
        },
    }

# ============================================================
# ROOT / HEALTH / STATUS
# ============================================================

@app.get("/")
async def root():
    return {
        "name": "AI Interview Platform + ATS Analyzer API",
        "version": "2.5.0",
        "status": "operational",
        "engines": {
            "llm":        LLM_CLIENT    is not None,
            "audio":      AUDIO_ENGINE  is not None,
            "video":      VIDEO_ENGINE  is not None,
            "video_version": VIDEO_VERSION,
            "evaluation": EVAL_FN       is not None,
            "questions":  QUESTION_FN   is not None,
            "stt":        STT_ENGINE    is not None,
            "tts":        TTS_FN        is not None,
            "groq_ats":   groq_client   is not None,
        },
        "cv_parsers":      {"pdf": _PYPDF_OK, "docx": _DOCX_OK},
        "active_sessions": len(sessions),
        "ats_endpoint":    "/ats/match",
    }


@app.get("/health")
async def health():
    return {
        "status":     "healthy",
        "timestamp":  datetime.utcnow().isoformat(),
        "engines": {
            "llm":        LLM_CLIENT   is not None,
            "audio":      AUDIO_ENGINE is not None,
            "video":      VIDEO_ENGINE is not None,
            "evaluation": EVAL_FN      is not None,
            "questions":  QUESTION_FN  is not None,
            "stt":        STT_ENGINE   is not None,
            "groq_ats":   groq_client  is not None,
        },
        "cv_parsers": {"pdf": _PYPDF_OK, "docx": _DOCX_OK},
    }


@app.get("/api/engines/status")
async def engine_status():
    return {
        "llm":        {"loaded": LLM_CLIENT   is not None, "provider": "Groq"},
        "audio":      {"loaded": AUDIO_ENGINE is not None},
        "video":      {"loaded": VIDEO_ENGINE is not None, "version": VIDEO_VERSION},
        "evaluation": {"loaded": EVAL_FN      is not None},
        "questions":  {"loaded": QUESTION_FN  is not None},
        "stt":        {"loaded": STT_ENGINE   is not None},
        "tts":        {"loaded": TTS_FN       is not None},
        "groq_ats":   {"loaded": groq_client  is not None},
        "cv_parsers": {"pdf": _PYPDF_OK, "docx": _DOCX_OK},
    }

# ============================================================
# CV PARSE
# ============================================================

@app.post("/api/cv/parse")
async def parse_cv(file: UploadFile = File(...)):
    """Parse a CV file (PDF / DOCX / DOC / TXT) and return the extracted text."""
    allowed = {".pdf", ".docx", ".doc", ".txt"}
    ext     = os.path.splitext((file.filename or "").lower())[1]
    if ext not in allowed:
        raise HTTPException(400, f"Unsupported file type '{ext}'. Allowed: PDF, DOCX, DOC, TXT")

    file_bytes = await file.read()
    cv_text    = extract_cv_text(file_bytes, file.filename or "cv.pdf")
    if cv_text.startswith("[CV_PARSE_ERROR]"):
        raise HTTPException(422, cv_text)

    return {
        "success":    True,
        "filename":   file.filename,
        "format":     ext.lstrip("."),
        "char_count": len(cv_text),
        "cv_text":    cv_text,
    }

# ============================================================
# ★  ATS ANALYZER ENDPOINTS
# ============================================================

@app.post("/ats/match")
async def ats_match(
    cv_file: UploadFile = File(...),
    jd_text: str        = Form(...),
):
    """
    Upload a CV (PDF/DOCX/TXT) + job description text.
    Returns ATS score, matched/missing keywords, and AI-generated improvement feedback.
    """
    ext = os.path.splitext((cv_file.filename or "").lower())[1]
    if ext not in {".pdf", ".docx", ".doc", ".txt", ".md"}:
        raise HTTPException(400, f"Unsupported file type '{ext}'.")

    file_bytes = await cv_file.read()
    cv_text    = extract_cv_text(file_bytes, cv_file.filename or "cv.pdf")

    if cv_text.startswith("[CV_PARSE_ERROR]"):
        raise HTTPException(422, cv_text)

    ats, matched, missing = _ats_score(cv_text, jd_text)
    analysis              = _llm_feedback(cv_text, jd_text, ats)

    return {
        "ats_score":        ats,
        "matched_keywords": matched,
        "missing_keywords": missing,
        "analysis":         analysis,
    }


@app.post("/ats/match-text")
async def ats_match_text(
    cv_text: str = Form(...),
    jd_text: str = Form(...),
):
    """
    Same as /ats/match but accepts CV as plain text (no file upload needed).
    """
    ats, matched, missing = _ats_score(cv_text, jd_text)
    analysis              = _llm_feedback(cv_text, jd_text, ats)

    return {
        "ats_score":        ats,
        "matched_keywords": matched,
        "missing_keywords": missing,
        "analysis":         analysis,
    }

# ============================================================
# ★  SESSION START
# ============================================================

@app.post("/api/session/start", response_model=SessionStartResponse)
async def start_session(
    job_title:       str                    = Form(...),
    job_description: str                    = Form(...),
    cv_text:         Optional[str]          = Form(None),
    cv_file:         Optional[UploadFile]   = File(None),
):
    if not job_title.strip():
        raise HTTPException(400, "job_title is required")

    final_cv, cv_parsed = "", False
    if cv_file and cv_file.filename:
        raw    = await cv_file.read()
        parsed = extract_cv_text(raw, cv_file.filename)
        if parsed.startswith("[CV_PARSE_ERROR]"):
            raise HTTPException(422, parsed)
        final_cv, cv_parsed = parsed, True
        print(f"📄 CV parsed: '{cv_file.filename}' ({len(final_cv)} chars)")
    elif cv_text:
        final_cv = cv_text.strip()
        print(f"📝 CV as text ({len(final_cv)} chars)")

    session    = create_session(job_title.strip(), job_description.strip(), final_cv)
    session_id = session["id"]
    sessions[session_id] = session

    first_stage = STAGE_ORDER[0]
    print(f"✅ Session created: {session_id} | Job: {job_title}")

    return SessionStartResponse(
        session_id=session_id,
        stage=first_stage,
        stage_index=0,
        question_number=1,
        total_questions=STAGE_QUESTION_COUNT[first_stage],
        cv_parsed=cv_parsed,
        status="active",
        message="Session created. Call /next-question to get the first question.",
    )

# ============================================================
# ★  NEXT QUESTION
# ============================================================

@app.get("/api/session/{session_id}/next-question", response_model=NextQuestionResponse)
async def next_question(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")

    session = sessions[session_id]
    if session["status"] != "active":
        raise HTTPException(409, "Interview is already completed")

    stage      = session["current_stage"]
    s_idx      = session["current_stage_index"]
    q_idx      = session["current_question_index"]
    max_q      = STAGE_QUESTION_COUNT[stage]
    stage_data = session["stages"][stage]

    if q_idx >= max_q:
        raise HTTPException(409, "All questions for this stage answered. Submit answer to advance.")

    if q_idx < len(stage_data["questions"]):
        question = stage_data["questions"][q_idx]
    else:
        question = _generate_adaptive_question(session, stage)
        stage_data["questions"].append(question)

    is_first  = (q_idx == 0)
    audio_b64 = TTS_FN(question) if TTS_FN else None

    print(f"📢 [{stage.upper()} Q{q_idx+1}/{max_q}] {question[:80]}...")

    return NextQuestionResponse(
        question=question,
        question_audio=audio_b64,
        stage=stage,
        stage_index=s_idx,
        question_number=q_idx + 1,
        total_questions=max_q,
        is_first=is_first,
        message=f"{stage.capitalize()} — Question {q_idx + 1} of {max_q}",
    )

# ============================================================
# ★  SESSION STATE
# ============================================================

@app.get("/api/session/{session_id}/state")
async def get_session_state(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    return _get_current_state(sessions[session_id])

# ============================================================
# ANSWER SUBMIT
# ============================================================

@app.post("/api/session/{session_id}/answer", response_model=AnswerSubmitResponse)
async def submit_answer(session_id: str, request: AnswerSubmitRequest):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")

    session = sessions[session_id]
    if session["status"] != "active":
        raise HTTPException(400, "Session is not active")

    current_stage = session["current_stage"]
    stage_data    = session["stages"][current_stage]
    current_q_idx = session["current_question_index"]

    if current_q_idx >= len(stage_data["questions"]):
        raise HTTPException(400, "No current question found. Call /next-question first.")

    current_question = stage_data["questions"][current_q_idx]

    evaluation = _evaluate_answer_safe(
        question=current_question,
        answer=request.answer,
        audio_features=request.audio_features,
        video_features=request.video_features,
        job_title=session["job_title"],
    )

    stage_data["answers"].append(request.answer)
    stage_data["evaluations"].append(evaluation)
    if request.audio_features:
        stage_data["audio_features"].append(request.audio_features)
    if request.video_features:
        stage_data["video_features"].append(request.video_features)

    session["current_question_index"] += 1
    session["global_question_index"]  += 1
    next_q_idx    = session["current_question_index"]
    max_questions = STAGE_QUESTION_COUNT[current_stage]

    stage_complete, interview_complete, next_stage = False, False, None
    message = f"Answer recorded for Q{current_q_idx + 1}."

    if next_q_idx >= max_questions:
        stage_complete = True
        next_stage_idx = session["current_stage_index"] + 1

        if next_stage_idx < len(STAGE_ORDER):
            next_stage = STAGE_ORDER[next_stage_idx]
            session["current_stage_index"]   = next_stage_idx
            session["current_stage"]          = next_stage
            session["current_question_index"] = 0
            message = f"{current_stage.capitalize()} complete! Moving to {next_stage.capitalize()}. Call /next-question."
        else:
            interview_complete = True
            session["status"]  = "completed"
            message = "Interview complete! Call /feedback for your results."

    eval_summary = {
        "technical_score":     evaluation.get("technical_score",    0),
        "communication_score": evaluation.get("communication_score",0),
        "confidence_index":    evaluation.get("confidence_index",   0),
        "short_feedback":      evaluation.get("short_feedback",      ""),
        "recommendation":      evaluation.get("overall_recommendation",""),
    }

    next_question_text  = None
    next_question_audio = None

    if not interview_complete:
        try:
            new_stage      = session["current_stage"]
            new_stage_data = session["stages"][new_stage]
            new_q_idx      = session["current_question_index"]

            if new_q_idx < len(new_stage_data["questions"]):
                next_question_text = new_stage_data["questions"][new_q_idx]
            else:
                next_question_text = _generate_adaptive_question(session, new_stage)
                new_stage_data["questions"].append(next_question_text)

            next_question_audio = TTS_FN(next_question_text) if TTS_FN else None
            print(f"📢 [{new_stage.upper()} Q{new_q_idx+1}/{STAGE_QUESTION_COUNT[new_stage]}] {next_question_text[:80]}...")
        except Exception as e:
            print(f"⚠️  Could not pre-fetch next question: {e}")

    return AnswerSubmitResponse(
        success=True,
        evaluation_summary=eval_summary,
        stage_complete=stage_complete,
        interview_complete=interview_complete,
        next_stage=next_stage,
        next_question=next_question_text,
        next_question_audio=next_question_audio,
        message=message,
    )

# ============================================================
# STATUS / DELETE
# ============================================================

@app.get("/api/session/{session_id}/status")
async def get_session_status(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    return _get_current_state(sessions[session_id])


@app.delete("/api/session/{session_id}")
async def end_session(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    sessions.pop(session_id, None)
    return {"success": True, "message": "Session ended"}

# ============================================================
# FINAL FEEDBACK
# ============================================================

@app.get("/api/session/{session_id}/feedback", response_model=FinalFeedbackResponse)
async def get_final_feedback(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")

    session = sessions[session_id]
    result  = _compute_final_feedback(session)

    return FinalFeedbackResponse(
        success=True,
        overall_score=          result["overall_score"],
        technical_score=        result["technical_score"],
        communication_score=    result["communication_score"],
        confidence_index=       result["confidence_index"],
        stress_index=           result["stress_index"],
        professional_presence=  result["professional_presence"],
        strengths=              result["strengths"],
        weaknesses=             result["weaknesses"],
        improvement_suggestions=result["improvement_suggestions"],
        stage_summaries=        result["stage_summaries"],
        recommendation=         result["recommendation"],
        message="Interview analysis complete.",
    )

# ============================================================
# ANALYSIS PAGE
# ============================================================

@app.get("/api/session/{session_id}/analysis/communication")
async def get_communication_analysis(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    session = sessions[session_id]
    audio   = _aggregate_audio(session)
    video   = _aggregate_video(session)
    if not audio:
        return {"success": True, "has_data": False}
    result = _score_communication(audio, video)
    return {"success": True, "has_data": True, **result}


@app.get("/api/session/{session_id}/analysis/body-language")
async def get_body_language_analysis(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    session = sessions[session_id]
    video   = _aggregate_video(session)
    if not video:
        return {"success": True, "has_data": False}
    result = _score_body_language(video)
    return {"success": True, "has_data": True, **result}


@app.get("/api/session/{session_id}/analysis/emotions")
async def get_emotion_analysis(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    session    = sessions[session_id]
    video      = _aggregate_video(session)
    per_stage  = {}

    for stage in STAGE_ORDER:
        sv = session["stages"][stage]["video_features"]
        if sv:
            def vn(key, subkey):
                vals = [v.get(key, {}).get(subkey)
                        for v in sv
                        if isinstance(v.get(key, {}).get(subkey), (int, float))]
                return round(sum(vals) / len(vals), 1) if vals else 0
            per_stage[stage] = {
                "confidence":  vn("behavioral", "confidence_score"),
                "stress":      vn("behavioral", "stress_level"),
                "eye_contact": vn("gaze",       "eye_contact_percentage"),
                "posture":     vn("posture",     "upright_percentage"),
                "questions":   len(sv),
            }

    all_alerts, emotion_timeline = [], []
    for stage in STAGE_ORDER:
        for i, vf in enumerate(session["stages"][stage]["video_features"]):
            all_alerts.extend(vf.get("alerts", []))
            emotion_timeline.append({
                "stage":            stage,
                "question":         i + 1,
                "dominant_emotion": vf.get("emotions", {}).get("dominant_emotion", "neutral"),
                "positive_ratio":   vf.get("emotions", {}).get("positive_ratio",   0),
                "negative_ratio":   vf.get("emotions", {}).get("negative_ratio",   0),
            })

    return {
        "success":          True,
        "has_data":         bool(video),
        "overall_score":    _score_body_language(video).get("overall_score", 0) if video else 0,
        "metrics":          _score_body_language(video).get("metrics", {})      if video else {},
        "per_stage":        per_stage,
        "emotion_breakdown":video.get("emotions", {}),
        "emotion_timeline": emotion_timeline,
        "gaze_data":        video.get("gaze", {}),
        "alerts":           all_alerts,
    }

# ============================================================
# QUICK ANALYSIS & REAL-TIME VIDEO
# ============================================================

@app.post("/api/analyze/quick")
async def quick_analyze(request: QuickAnalysisRequest):
    evaluation = _evaluate_answer_safe(
        question=request.question,
        answer=request.answer,
        audio_features=request.audio_features,
        video_features=request.video_features,
        job_title="General",
    )
    return {"success": True, "evaluation": evaluation}


@app.post("/api/session/{session_id}/video-frame")
async def submit_video_frame(session_id: str, request: VideoFrameRequest):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    buf = sessions[session_id]["video_frames_buffer"]
    buf.append({"frame": request.frame_base64, "timestamp": time.time()})
    if len(buf) > 150:
        sessions[session_id]["video_frames_buffer"] = buf[-150:]
    return {"success": True, "buffered_frames": len(sessions[session_id]["video_frames_buffer"])}


@app.get("/api/session/{session_id}/video-analysis")
async def get_video_analysis(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    buf = sessions[session_id]["video_frames_buffer"]
    if not buf or not VIDEO_ENGINE:
        return {"success": True, "has_data": False, "video_version": VIDEO_VERSION}
    try:
        import base64, numpy as np, cv2
        frames = []
        for item in buf[-60:]:
            nparr = np.frombuffer(base64.b64decode(item["frame"]), np.uint8)
            frame = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            if frame is not None:
                frames.append(frame)
        if frames:
            features = VIDEO_ENGINE.extract_features_from_video(frames, len(frames) / 30.0)
            return {
                "success":        True,
                "has_data":       True,
                "features":       features,
                "frames_analyzed":len(frames),
                "video_version":  VIDEO_VERSION,
            }
    except Exception as e:
        print(f"⚠️  Video analysis error: {e}")
    return {"success": False, "message": "Video analysis failed", "video_version": VIDEO_VERSION}

# ============================================================
# FILE UPLOAD ENDPOINTS
# ============================================================

@app.post("/api/session/{session_id}/audio-analysis")
async def analyze_audio(session_id: str, file: UploadFile = File(...)):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    if not AUDIO_ENGINE:
        raise HTTPException(500, "Audio Engine not loaded")
    try:
        audio_bytes   = await file.read()
        transcription = ""

        if STT_ENGINE:
            header = audio_bytes[:16]
            if   header[:4] == b"\x1a\x45\xdf\xa3": real_suffix = ".webm"
            elif header[:4] == b"OggS":              real_suffix = ".ogg"
            elif header[:4] == b"RIFF" and header[8:12] == b"WAVE": real_suffix = ".wav"
            elif header[:3] == b"ID3" or header[:2] == b"\xff\xfb": real_suffix = ".mp3"
            else: real_suffix = os.path.splitext(file.filename or ".wav")[1] or ".wav"

            with tempfile.NamedTemporaryFile(delete=False, suffix=real_suffix) as tmp:
                tmp.write(audio_bytes)
                tmp_path = tmp.name
            try:
                transcription = STT_ENGINE.transcribe_file(tmp_path)
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)

        try:
            audio_data, sr = sf.read(io.BytesIO(audio_bytes))
        except Exception:
            try:
                import librosa
                header  = audio_bytes[:4]
                tmp_ext = ".webm" if header == b"\x1a\x45\xdf\xa3" else (".ogg" if header == b"OggS" else ".webm")
                with tempfile.NamedTemporaryFile(delete=False, suffix=tmp_ext) as tf:
                    tf.write(audio_bytes)
                    tf_path = tf.name
                try:
                    audio_data, sr = librosa.load(tf_path, sr=None, mono=False)
                    if audio_data.ndim > 1:
                        audio_data = audio_data.mean(axis=0)
                finally:
                    if os.path.exists(tf_path):
                        os.remove(tf_path)
            except Exception as fallback_err:
                print(f"⚠️  Audio decode fallback failed: {fallback_err}")
                return {"success": True, "audio_features": {}, "transcription": transcription, "sample_rate": 16000}

        if audio_data.ndim > 1:
            audio_data = audio_data.mean(axis=1)
        features = AUDIO_ENGINE.extract_features(audio_data)

        return {"success": True, "audio_features": features, "transcription": transcription, "sample_rate": sr}
    except Exception as e:
        raise HTTPException(400, f"Failed to process audio: {e}")


@app.post("/api/session/{session_id}/test-video-upload")
async def test_video_upload(session_id: str, file: UploadFile = File(...)):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    if not VIDEO_ENGINE:
        raise HTTPException(500, "Video Engine not loaded")
    try:
        import cv2, numpy as np
        suffix = os.path.splitext(file.filename or ".mp4")[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        frames, cap = [], cv2.VideoCapture(tmp_path)
        fps      = cap.get(cv2.CAP_PROP_FPS) or 30
        interval = max(1, int(fps))
        count    = 0
        while cap.isOpened() and len(frames) < 30:
            ret, frame = cap.read()
            if not ret:
                break
            if count % interval == 0:
                frames.append(cv2.resize(frame, (640, 480)))
            count += 1
        cap.release()
        os.remove(tmp_path)

        if not frames:
            raise ValueError("No frames extracted")

        features = VIDEO_ENGINE.extract_features_from_video(frames, float(len(frames)))
        return {"success": True, "message": f"Analyzed {len(frames)} frames", "video_features": features}
    except Exception as e:
        raise HTTPException(400, f"Failed to process video: {e}")

# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    import uvicorn
    print("\n🚀 AI Interview Platform + ATS Analyzer v2.5 starting...")
    print("   URL:       http://localhost:8000")
    print("   Docs:      http://localhost:8000/docs")
    print("   ATS Match: http://localhost:8000/ats/match\n")
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")
