"""
CV Parser Utility
=================
Extracts plain text from uploaded CV files.

Supported formats:
  - PDF  (.pdf)   — via pypdf
  - Word (.docx)  — via python-docx
  - Word (.doc)   — via pandoc (legacy format)
  - Plain text    — fallback

Usage (standalone):
    from app.chat_layer.utils.cv_parser import extract_cv_text
    text = extract_cv_text(file_bytes, filename="cv.pdf")

Usage (FastAPI upload):
    text = await parse_cv_upload(file)   # UploadFile
"""

import io
import os
import tempfile
from typing import Optional


# ── availability flags ──────────────────────────────────────────────────────

_PYPDF_OK = False
_DOCX_OK  = False

try:
    from pypdf import PdfReader
    _PYPDF_OK = True
except ImportError:
    try:
        from PyPDF2 import PdfReader          # older alias
        _PYPDF_OK = True
    except ImportError:
        pass

try:
    from docx import Document as DocxDocument
    _DOCX_OK = True
except ImportError:
    pass


# ── public API ───────────────────────────────────────────────────────────────

def extract_cv_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from a CV file.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename:   Original filename (used to detect format).

    Returns:
        Extracted text as a single string, or an error message prefixed
        with "[CV_PARSE_ERROR]" so the caller can decide what to do.
    """
    ext = os.path.splitext(filename.lower())[1]

    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext == ".docx":
        return _parse_docx(file_bytes)
    elif ext == ".doc":
        return _parse_doc_legacy(file_bytes)
    elif ext in (".txt", ".md", ".rtf"):
        # best-effort UTF-8 / latin-1 decode
        return _decode_text(file_bytes)
    else:
        return f"[CV_PARSE_ERROR] Unsupported file type: '{ext}'. Please upload PDF or DOCX."


async def parse_cv_upload(upload_file) -> str:
    """
    Convenience wrapper for a FastAPI UploadFile object.

    Args:
        upload_file: fastapi.UploadFile

    Returns:
        Extracted CV text string.
    """
    file_bytes = await upload_file.read()
    return extract_cv_text(file_bytes, upload_file.filename or "cv.pdf")


# ── format parsers ───────────────────────────────────────────────────────────

def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pypdf."""
    if not _PYPDF_OK:
        return "[CV_PARSE_ERROR] pypdf is not installed. Run: pip install pypdf"

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages  = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())

        full_text = "\n\n".join(pages).strip()

        if not full_text:
            return "[CV_PARSE_ERROR] PDF appears to be scanned (no extractable text). Please upload a text-based PDF or DOCX."

        return full_text

    except Exception as exc:
        return f"[CV_PARSE_ERROR] Failed to parse PDF: {exc}"


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file using python-docx."""
    if not _DOCX_OK:
        return "[CV_PARSE_ERROR] python-docx is not installed. Run: pip install python-docx"

    try:
        doc   = DocxDocument(io.BytesIO(file_bytes))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also pull text from tables (sometimes CVs use table layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in lines:
                        lines.append(cell_text)

        full_text = "\n".join(lines).strip()

        if not full_text:
            return "[CV_PARSE_ERROR] DOCX file appears to be empty."

        return full_text

    except Exception as exc:
        return f"[CV_PARSE_ERROR] Failed to parse DOCX: {exc}"


def _parse_doc_legacy(file_bytes: bytes) -> str:
    """
    Extract text from a legacy .doc file via pandoc.
    Requires pandoc to be installed on the server.
    """
    try:
        import subprocess

        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        result = subprocess.run(
            ["pandoc", tmp_path, "-t", "plain"],
            capture_output=True,
            text=True,
            timeout=15,
        )
        os.remove(tmp_path)

        if result.returncode != 0:
            return f"[CV_PARSE_ERROR] pandoc failed: {result.stderr.strip()}"

        full_text = result.stdout.strip()
        if not full_text:
            return "[CV_PARSE_ERROR] Legacy .doc file produced no text."

        return full_text

    except FileNotFoundError:
        return "[CV_PARSE_ERROR] pandoc is not installed. For .doc files, please convert to .docx or PDF first."
    except Exception as exc:
        return f"[CV_PARSE_ERROR] Failed to parse .doc: {exc}"


def _decode_text(file_bytes: bytes) -> str:
    """Best-effort decode for plain text / markdown / RTF files."""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return "[CV_PARSE_ERROR] Could not decode text file."